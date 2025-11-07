"""Document generation service for creating synthetic documents"""

from typing import Dict, Any, Optional, List
from pathlib import Path
import json
from datetime import datetime
from faker import Faker
from docx import Document
from docx.shared import Inches, Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

from app.core.config import settings


class DocumentGenerationService:
    """
    Service for generating synthetic documents.

    Supports invoices, contracts, reports, and custom documents
    in multiple formats (PDF, DOCX, JSON).
    """

    def __init__(self, base_dir: Optional[Path] = None):
        """
        Initialize document generator.

        Args:
            base_dir: Base directory for document artifacts
        """
        if base_dir is None:
            base_dir = Path(settings.LOCAL_ARTIFACTS_DIR)
        self.base_dir = base_dir
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.faker = Faker()

    async def generate_invoice(
        self, job_id: str, data: Optional[Dict[str, Any]] = None, format: str = "pdf"
    ) -> Path:
        """
        Generate a synthetic invoice.

        Args:
            job_id: Job identifier
            data: Optional invoice data
            format: Output format (pdf, docx, json)

        Returns:
            Path to generated document
        """
        # Generate fake invoice data if not provided
        if data is None:
            data = self._generate_invoice_data()

        job_dir = self.base_dir / job_id
        job_dir.mkdir(parents=True, exist_ok=True)

        if format == "pdf":
            return self._generate_invoice_pdf(job_dir, data)
        elif format == "docx":
            return self._generate_invoice_docx(job_dir, data)
        else:  # json
            return self._generate_invoice_json(job_dir, data)

    def _generate_invoice_data(self) -> Dict[str, Any]:
        """Generate fake invoice data"""
        return {
            "invoice_number": f"INV-{self.faker.random_int(10000, 99999)}",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "due_date": self.faker.future_date(end_date="+30d").strftime("%Y-%m-%d"),
            "bill_to": {
                "name": self.faker.company(),
                "address": self.faker.address(),
                "email": self.faker.company_email(),
            },
            "items": [
                {
                    "description": self.faker.catch_phrase(),
                    "quantity": self.faker.random_int(1, 10),
                    "unit_price": round(self.faker.random.uniform(10, 500), 2),
                }
                for _ in range(self.faker.random_int(2, 5))
            ],
            "tax_rate": 0.08,
            "notes": "Payment due within 30 days",
        }

    def _generate_invoice_pdf(self, job_dir: Path, data: Dict[str, Any]) -> Path:
        """Generate invoice as PDF"""
        file_path = job_dir / "invoice.pdf"

        c = canvas.Canvas(str(file_path), pagesize=letter)
        width, height = letter

        # Title
        c.setFont("Helvetica-Bold", 24)
        c.drawString(50, height - 50, "INVOICE")

        # Invoice details
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 100, f"Invoice Number: {data['invoice_number']}")
        c.drawString(50, height - 120, f"Date: {data['date']}")
        c.drawString(50, height - 140, f"Due Date: {data['due_date']}")

        # Bill to
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, height - 180, "Bill To:")
        c.setFont("Helvetica", 11)
        c.drawString(50, height - 200, data["bill_to"]["name"])

        # Items
        y_position = height - 250
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y_position, "Items:")

        y_position -= 25
        subtotal = 0
        for item in data["items"]:
            total = item["quantity"] * item["unit_price"]
            subtotal += total
            c.setFont("Helvetica", 10)
            c.drawString(50, y_position, f"{item['description']}")
            c.drawString(350, y_position, f"Qty: {item['quantity']}")
            c.drawString(450, y_position, f"${total:.2f}")
            y_position -= 20

        # Totals
        y_position -= 20
        tax = subtotal * data.get("tax_rate", 0)
        total = subtotal + tax

        c.setFont("Helvetica-Bold", 11)
        c.drawString(350, y_position, "Subtotal:")
        c.drawString(450, y_position, f"${subtotal:.2f}")
        y_position -= 20
        c.drawString(350, y_position, f"Tax ({data.get('tax_rate', 0)*100}%):")
        c.drawString(450, y_position, f"${tax:.2f}")
        y_position -= 20
        c.drawString(350, y_position, "Total:")
        c.drawString(450, y_position, f"${total:.2f}")

        c.save()
        return file_path

    def _generate_invoice_docx(self, job_dir: Path, data: Dict[str, Any]) -> Path:
        """Generate invoice as DOCX"""
        file_path = job_dir / "invoice.docx"

        doc = Document()
        doc.add_heading("INVOICE", 0)

        doc.add_paragraph(f"Invoice Number: {data['invoice_number']}")
        doc.add_paragraph(f"Date: {data['date']}")
        doc.add_paragraph(f"Due Date: {data['due_date']}")

        doc.add_heading("Bill To:", level=2)
        doc.add_paragraph(data["bill_to"]["name"])

        doc.add_heading("Items:", level=2)

        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Description"
        hdr_cells[1].text = "Quantity"
        hdr_cells[2].text = "Unit Price"
        hdr_cells[3].text = "Total"

        subtotal = 0
        for item in data["items"]:
            row_cells = table.add_row().cells
            total = item["quantity"] * item["unit_price"]
            subtotal += total
            row_cells[0].text = item["description"]
            row_cells[1].text = str(item["quantity"])
            row_cells[2].text = f"${item['unit_price']:.2f}"
            row_cells[3].text = f"${total:.2f}"

        doc.add_paragraph()
        tax = subtotal * data.get("tax_rate", 0)
        total = subtotal + tax

        doc.add_paragraph(f"Subtotal: ${subtotal:.2f}")
        doc.add_paragraph(f"Tax: ${tax:.2f}")
        doc.add_paragraph(f"Total: ${total:.2f}")

        doc.save(str(file_path))
        return file_path

    def _generate_invoice_json(self, job_dir: Path, data: Dict[str, Any]) -> Path:
        """Generate invoice as JSON"""
        file_path = job_dir / "invoice.json"

        # Calculate totals
        subtotal = sum(item["quantity"] * item["unit_price"] for item in data["items"])
        tax = subtotal * data.get("tax_rate", 0)
        total = subtotal + tax

        data["subtotal"] = round(subtotal, 2)
        data["tax"] = round(tax, 2)
        data["total"] = round(total, 2)

        with open(file_path, "w") as f:
            json.dump(data, f, indent=2)

        return file_path


# Global instance
_doc_service: Optional[DocumentGenerationService] = None


def get_document_service() -> DocumentGenerationService:
    """Get or create the global document service instance"""
    global _doc_service
    if _doc_service is None:
        _doc_service = DocumentGenerationService()
    return _doc_service
